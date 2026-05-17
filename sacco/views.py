from django.shortcuts import render, redirect, get_object_or_404
from django.core.exceptions import ObjectDoesNotExist
from django.db import IntegrityError
from decimal import Decimal

from django.contrib import messages
from django.db.models import Sum, Count, Q
from django.utils import timezone
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
import csv
import json

import logging
from django.core.mail import send_mail
from django.conf import settings
import random
import string
from datetime import timedelta
from io import BytesIO

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from openpyxl import Workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# pandas import commented out - not used in loan functionality
# try:
#     import pandas as pd
#     PANDAS_AVAILABLE = True
# except ImportError:
#     PANDAS_AVAILABLE = False

from .models import Member, Savings, Loan, Payment, Withdrawal, OTP, Notification, NotificationReply
from .forms import MemberForm, SavingsForm, LoanForm, LoanApplicationForm, PaymentForm, WithdrawalForm, WithdrawalApplicationForm, ForgotPasswordForm, OTPVerificationForm, SetPasswordForm, NotificationForm
from .decorators import user_required, admin_required, superadmin_required, get_user_role
from decimal import Decimal

# Initialize logger
logger = logging.getLogger(__name__)

def ensure_member(user):
    """Ensure member exists for user, create basic if missing"""
    try:
        member = Member.objects.get(user=user)
    except Member.DoesNotExist:
        member = Member.objects.create(
            user=user,
            name=user.get_full_name() or user.username,
            phone=user.username,  # phone as username
            id_number=f"ID{user.id}",
            role='user',
            email=getattr(user, 'email', ''),
        )
        logger.info(f"Auto-created member {member.phone} for user {user.username}")
    return member

def send_sms(phone, message):
    """Placeholder SMS sender; replace with a real SMS gateway call in production."""
    print(f"[send_sms] To: {phone}; Message: {message}")
    return True

@user_required
def my_withdrawals(request):
    """Show user's withdrawal history"""
    try:
        member = Member.objects.get(user=request.user)
    except Member.DoesNotExist:
        messages.error(request, 'You are not a registered member.')
        return redirect('sacco:dashboard')

    withdrawals = Withdrawal.objects.filter(member=member).order_by('-application_date')
    return render(request, 'sacco/my_withdrawals.html', {'withdrawals': withdrawals})


def process_withdrawal_disbursement(withdrawal):
    """
    Process the actual disbursement of approved withdrawal
    This function can be called by admin or automated system
    """
    if withdrawal.status != 'approved':
        return False, "Withdrawal is not approved"

    # Mark as processed (you can add a processed field to the model if needed)
    # For now, we just ensure the savings transaction exists
    existing_transaction = Savings.objects.filter(
        member=withdrawal.member,
        amount=withdrawal.amount,
        transaction_type='withdrawal',
        transaction_date=withdrawal.approval_date
    ).exists()

    if not existing_transaction:
        # Create the transaction if it doesn't exist
        Savings.objects.create(
            member=withdrawal.member,
            amount=withdrawal.amount,
            transaction_type='withdrawal',
            payment_method=withdrawal.payment_method,
            transaction_date=withdrawal.approval_date or timezone.now().date()
        )

    # Here you can add integration with payment gateways, mobile money APIs, etc.
    # For example:
    # if withdrawal.payment_method == 'mobile_money':
    #     initiate_mobile_money_transfer(withdrawal.member.phone, withdrawal.amount)
    # elif withdrawal.payment_method == 'bank_transfer':
    #     initiate_bank_transfer(withdrawal.member.account_number, withdrawal.amount)

    return True, "Withdrawal processed successfully"


@login_required
def members_list(request):
    """List all members (staff/superuser only)"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')
    members = Member.objects.all().order_by('name')
    return render(request, 'sacco/members_list.html', {'members': members})

@login_required
def add_member(request):
    """Add a new member (role assignment only for superadmin)"""
    is_superuser = request.user.is_superuser
    
    if request.method == 'POST':
        form = MemberForm(request.POST)
        if form.is_valid():
            member = form.save()
            password = form.cleaned_data.get('password')
            
            # Get role (superadmin can assign roles, others default to 'user')
            if is_superuser:
                role = form.cleaned_data.get('role')
            else:
                role = 'user'
                member.role = 'user'
                member.save()
            
            # Create user account if phone and password are provided
            if member.phone and password:
                user = User.objects.create_user(
                    username=member.phone,  # Use phone as username
                    email=member.email or '',
                    first_name=member.name.split()[0] if member.name else '',
                    last_name=' '.join(member.name.split()[1:]) if member.name and len(member.name.split()) > 1 else ''
                )
                user.set_password(password)
                
                # Set permissions based on role (only if superadmin assigned a role)
                if is_superuser and role in ['admin', 'superadmin']:
                    if role == 'admin':
                        user.is_staff = True
                        user.is_superuser = False
                    elif role == 'superadmin':
                        user.is_staff = True
                        user.is_superuser = True
                else:
                    # Default to regular user
                    user.is_staff = False
                    user.is_superuser = False
                
                user.save()
                member.user = user
                member.save()
                
                role_text = f"with role: {role}" if is_superuser else "as regular user"
                logger.info(f"New member created: {member.name} {role_text} by {request.user.username}")
                messages.success(request, f'Member {member.name} added successfully {role_text}!')
            else:
                logger.info(f"New member added without user account: {member.name}")
                messages.success(request, 'Member added successfully (no login account created)!')
                
            return redirect('sacco:members_list')
    else:
        form = MemberForm()
        # Remove role field from form if user is not superuser
        if not is_superuser:
            if 'role' in form.fields:
                del form.fields['role']

    context = {
        'form': form,
        'is_superuser': is_superuser,
    }
    return render(request, 'sacco/add_member.html', context)

@login_required
def add_loan(request):
    """Add a new loan"""
    if request.method == 'POST':
        form = LoanForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Loan added successfully!')
            return redirect('sacco:loans_list')
    else:
        form = LoanForm()

    return render(request, 'sacco/add_loan.html', {'form': form})

@login_required
def add_savings(request):
    """Add a new savings transaction"""
    if request.method == 'POST':
        form = SavingsForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Savings transaction added successfully!')
            return redirect('sacco:savings_list')
    else:
        form = SavingsForm()

    return render(request, 'sacco/add_savings.html', {'form': form})

@login_required
def edit_member(request, member_id):
    """Edit an existing member (role assignment only for superadmin)"""
    member = get_object_or_404(Member, pk=member_id)
    is_superuser = request.user.is_superuser
    
    if request.method == 'POST':
        form = MemberForm(request.POST, instance=member)
        if form.is_valid():
            old_role = member.role
            member = form.save()
            new_role = member.role
            
            # Reset role if non-superuser tried to change it (security measure)
            if not is_superuser and old_role != new_role:
                member.role = old_role
                member.save()
                messages.warning(request, 'You do not have permission to change member roles. Only superadmins can assign roles.')
                new_role = old_role
            
            # Update user permissions if role changed (only for superadmin changes)
            if is_superuser and member.user and old_role != new_role:
                if new_role == 'admin':
                    member.user.is_staff = True
                    member.user.is_superuser = False
                elif new_role == 'superadmin':
                    member.user.is_staff = True
                    member.user.is_superuser = True
                else:  # user
                    member.user.is_staff = False
                    member.user.is_superuser = False
                
                member.user.save()
                logger.info(f"Member role updated by {request.user.username}: {member.name} from {old_role} to {new_role}")
            
            messages.success(request, 'Member updated successfully!')
            return redirect('sacco:members_list')
    else:
        form = MemberForm(instance=member)
        # Remove role field from form if user is not superuser
        if not is_superuser:
            if 'role' in form.fields:
                del form.fields['role']

    context = {
        'form': form,
        'member': member,
        'is_superuser': is_superuser,
    }
    return render(request, 'sacco/edit_member.html', context)

@login_required
def delete_member(request, member_id):
    """Delete a member"""
    member = get_object_or_404(Member, pk=member_id)
    if request.method == 'POST':
        member.delete()
        messages.success(request, 'Member deleted successfully!')
        return redirect('sacco:members_list')
    return render(request, 'sacco/delete_member_confirm.html', {'member': member})

@login_required
def savings_list(request):
    # allow only staff/superuser to view lists (adjust as needed)
    if not (request.user.is_staff or request.user.is_superuser):
        return redirect('sacco:dashboard')

    # defensive: exclude any objects without a valid primary key
    savings = Savings.objects.exclude(pk__isnull=True).select_related('member').order_by('-transaction_date')
    return render(request, 'sacco/savings_list.html', {'savings': savings})

@login_required
def loans_list(request):
    # allow only staff/superuser to view lists (adjust as needed)
    if not (request.user.is_staff or request.user.is_superuser):
        return redirect('sacco:dashboard')

    # exclude objects with no valid primary key (defensive)
    loans = Loan.objects.exclude(pk__isnull=True).select_related('member').order_by('-application_date')
    return render(request, 'sacco/loans_list.html', {'loans': loans})

@login_required
def payments_list(request):
    """List all payments"""
    payments = Payment.objects.select_related('loan__member').all()
    return render(request, 'sacco/payments_list.html', {'payments': payments})

@login_required
def add_payment(request):
    """Add a new payment"""
    if request.method == 'POST':
        form = PaymentForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Payment added successfully!')
            return redirect('sacco:payments_list')
    else:
        form = PaymentForm()

    return render(request, 'sacco/add_payment.html', {'form': form})

@login_required
def edit_loan(request, loan_id):
    """Edit an existing loan (admin only)"""
    # Admin only
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')
    
    loan = get_object_or_404(Loan, pk=loan_id)
    old_status = loan.status
    
    if request.method == 'POST':
        form = LoanForm(request.POST, instance=loan)
        if form.is_valid():
            form.save()
            
            # Check if status changed and create notifications
            if loan.status != old_status:
                loan.approval_date = timezone.now().date()
                loan.save()
                
                if loan.status == 'approved':
                    # Create notification for member
                    Notification.objects.create(
                        member=loan.member,
                        notification_type='loan_approval',
                        title=f'Your Loan Application has been Approved',
                        message=f'Congratulations! Your loan application for UGX {loan.amount} at {loan.interest_rate}% interest for {loan.term_months} months has been APPROVED.',
                        loan=loan,
                        admin_user=request.user,
                    )
                    
                    # Create notification for admins
                    admin_members = Member.objects.filter(role__in=['admin', 'superadmin'])
                    for admin_member in admin_members:
                        if admin_member.user != request.user:
                            Notification.objects.create(
                                member=admin_member,
                                notification_type='loan_approval',
                                title=f'Loan Approved: {loan.member.name}',
                                message=f'Admin {request.user.first_name or request.user.username} has APPROVED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).',
                                loan=loan,
                                admin_user=request.user,
                            )
                    
                    messages.success(request, 'Loan updated and approved! Member notified.')
                
                elif loan.status == 'rejected':
                    # Create notification for member
                    Notification.objects.create(
                        member=loan.member,
                        notification_type='loan_rejection',
                        title=f'Your Loan Application has been Rejected',
                        message=f'Unfortunately, your loan application for UGX {loan.amount} has been REJECTED. Please contact admin for more details.',
                        loan=loan,
                        admin_user=request.user,
                    )
                    
                    # Create notification for admins
                    admin_members = Member.objects.filter(role__in=['admin', 'superadmin'])
                    for admin_member in admin_members:
                        if admin_member.user != request.user:
                            Notification.objects.create(
                                member=admin_member,
                                notification_type='loan_rejection',
                                title=f'Loan Rejected: {loan.member.name}',
                                message=f'Admin {request.user.first_name or request.user.username} has REJECTED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).',
                                loan=loan,
                                admin_user=request.user,
                            )
                    
                    messages.success(request, 'Loan updated and rejected! Member notified.')
                else:
                    messages.success(request, 'Loan updated successfully!')
            else:
                messages.success(request, 'Loan updated successfully!')
            
            return redirect('sacco:loans_list')
    else:
        form = LoanForm(instance=loan)

    return render(request, 'sacco/edit_loan.html', {'form': form, 'loan': loan})

@login_required
def create_loan_notification(loan, status, admin_user):
    """Helper function to create loan status notifications"""
    try:
        if status == 'approved':
            notification_type = 'loan_approval'
            member_title = 'Your Loan Application has been Approved'
            member_message = f'Congratulations! Your loan application for UGX {loan.amount} at {loan.interest_rate}% interest for {loan.term_months} months has been APPROVED.'
            admin_title = f'Loan Approved: {loan.member.name}'
            admin_message = f'Admin {admin_user.first_name or admin_user.username} has APPROVED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).'
        elif status == 'rejected':
            notification_type = 'loan_rejection'
            member_title = 'Your Loan Application has been Rejected'
            member_message = f'Unfortunately, your loan application for UGX {loan.amount} has been REJECTED. Please contact admin for more details.'
            admin_title = f'Loan Rejected: {loan.member.name}'
            admin_message = f'Admin {admin_user.first_name or admin_user.username} has REJECTED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).'
        else:
            return

        # Create notification for member
        Notification.objects.create(
            member=loan.member,
            notification_type=notification_type,
            title=member_title,
            message=member_message,
            loan=loan,
            admin_user=admin_user,
        )
        
        # Create notifications for other admins
        try:
            admin_members = Member.objects.filter(role__in=['admin', 'superadmin'])
            for admin_member in admin_members:
                if admin_member.user and admin_member.user != admin_user:
                    Notification.objects.create(
                        member=admin_member,
                        notification_type=notification_type,
                        title=admin_title,
                        message=admin_message,
                        loan=loan,
                        admin_user=admin_user,
                    )
        except Exception as e:
            logger.error(f"Error creating admin notifications for loan {loan.id}: {str(e)}")
            # Don't let this error stop the main notification
            
    except Exception as e:
        logger.error(f"Error creating loan notification for loan {loan.id}: {str(e)}")
        # Log but don't raise - we don't want notification errors to break the approval/rejection

@login_required
def approve_loan(request, loan_id):
    """Approve a loan (admin only)"""
    try:
        if not (request.user.is_staff or request.user.is_superuser):
            messages.error(request, "Permission denied.")
            return redirect('sacco:dashboard')
        
        loan = get_object_or_404(Loan, pk=loan_id)
        if loan.status == 'approved':
            messages.warning(request, 'Loan is already approved.')
            return redirect('sacco:loans_list')
        
        loan.status = 'approved'
        loan.approval_date = timezone.now().date()
        loan.save()
        create_loan_notification(loan, 'approved', request.user)
        messages.success(request, f'Loan #{loan.loan_id} approved successfully! Member notified.')
        logger.info(f"Loan #{loan.loan_id} approved by {request.user.username}")
    except Exception as e:
        logger.error(f"Error approving loan {loan_id}: {str(e)}")
        messages.error(request, f'An error occurred while approving the loan: {str(e)}')
    
    return redirect('sacco:loans_list')

@login_required
def reject_loan(request, loan_id):
    """Reject a loan (admin only)"""
    try:
        if not (request.user.is_staff or request.user.is_superuser):
            messages.error(request, "Permission denied.")
            return redirect('sacco:dashboard')
        
        loan = get_object_or_404(Loan, pk=loan_id)
        if loan.status == 'rejected':
            messages.warning(request, 'Loan is already rejected.')
            return redirect('sacco:loans_list')
        
        loan.status = 'rejected'
        loan.approval_date = timezone.now().date()
        loan.save()
        create_loan_notification(loan, 'rejected', request.user)
        messages.success(request, f'Loan #{loan.loan_id} rejected. Member notified.')
        logger.info(f"Loan #{loan.loan_id} rejected by {request.user.username}")
    except Exception as e:
        logger.error(f"Error rejecting loan {loan_id}: {str(e)}")
        messages.error(request, f'An error occurred while rejecting the loan: {str(e)}')
    
    return redirect('sacco:loans_list')

def delete_loan(request, loan_id):
    """Delete a loan (admin only)"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')
    
    loan = get_object_or_404(Loan, pk=loan_id)
    if request.method == 'POST':
        loan.delete()
        messages.success(request, 'Loan deleted successfully!')
        return redirect('sacco:loans_list')
    return render(request, 'sacco/delete_loan_confirm.html', {'loan': loan})

@login_required
def edit_savings(request, transaction_id):
    """Edit an existing savings transaction"""
    saving = get_object_or_404(Savings, pk=transaction_id)
    if request.method == 'POST':
        form = SavingsForm(request.POST, instance=saving)
        if form.is_valid():
            form.save()
            messages.success(request, 'Savings updated successfully!')
            return redirect('sacco:savings_list')
    else:
        form = SavingsForm(instance=saving)

    return render(request, 'sacco/edit_savings.html', {'form': form, 'saving': saving})

@login_required
def delete_savings(request, transaction_id):
    """Delete a savings transaction"""
    saving = get_object_or_404(Savings, pk=transaction_id)
    if request.method == 'POST':
        saving.delete()
        messages.success(request, 'Savings deleted successfully!')
        return redirect('sacco:savings_list')
    return render(request, 'sacco/delete_savings_confirm.html', {'saving': saving})

def reports(request):
    """Reports view: render a simple summary template."""
    total_members = Member.objects.count()
    total_savings = Savings.objects.filter(transaction_type='deposit').aggregate(total=Sum('amount'))['total'] or 0
    total_loans_given_out = Loan.objects.filter(status='approved').aggregate(total=Sum('amount'))['total'] or 0
    pending_loans = Loan.objects.filter(status='pending').count()
    pending_loans_amount = Loan.objects.filter(status='pending').aggregate(total=Sum('amount'))['total'] or 0
    total_loan_payments = Payment.objects.aggregate(total=Sum('amount'))['total'] or 0

    context = {
        'total_members': total_members,
        'total_savings': total_savings,
        'total_loans_given_out': total_loans_given_out,
        'pending_loans': pending_loans,
        'pending_loans_amount': pending_loans_amount,
        'total_loan_payments': total_loan_payments,
    }
    return render(request, 'sacco/reports.html', context)

def export_members_report(request):
    """
    Export members as CSV. Staff/superuser only.
    """
    if not (request.user.is_staff or request.user.is_superuser):
        return redirect('sacco:dashboard')

    # Force a safe ordering that exists on the model (use actual field name or pk)
    qs = Member.objects.all().order_by('member_id')  # use 'member_id' (or 'pk') instead of 'id'

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="members_report.csv"'

    writer = csv.writer(response)
    writer.writerow(['member_id', 'name', 'id_number', 'phone', 'email', 'join_date'])
    for m in qs:
        writer.writerow([
            getattr(m, 'member_id', '') or '',
            getattr(m, 'name', '') or '',
            getattr(m, 'id_number', '') or '',
            getattr(m, 'phone', '') or '',
            getattr(m, 'email', '') or '',
            (m.join_date.isoformat() if getattr(m, 'join_date', None) else ''),
        ])
    return response

def export_data(request, model_name, format_type):
    """
    Generic export function for all models supporting CSV, PDF, Excel, and DOCX formats
    """
    if not (request.user.is_staff or request.user.is_superuser):
        return redirect('sacco:dashboard')

    # Define model mappings
    model_configs = {
        'members': {
            'model': Member,
            'fields': ['member_id', 'name', 'id_number', 'phone', 'email', 'join_date', 'role'],
            'headers': ['Member ID', 'Name', 'ID Number', 'Phone', 'Email', 'Join Date', 'Role'],
            'title': 'Members Report'
        },
        'savings': {
            'model': Savings,
            'fields': ['transaction_id', 'member__name', 'amount', 'transaction_type', 'transaction_date', 'payment_method'],
            'headers': ['Transaction ID', 'Member', 'Amount', 'Type', 'Date', 'Payment Method'],
            'title': 'Savings Report'
        },
        'loans': {
            'model': Loan,
            'fields': ['loan_id', 'member__name', 'amount', 'interest_rate', 'term_months', 'status', 'application_date', 'approval_date'],
            'headers': ['Loan ID', 'Member', 'Amount', 'Interest Rate', 'Term (Months)', 'Status', 'Application Date', 'Approval Date'],
            'title': 'Loans Report'
        },
        'withdrawals': {
            'model': Withdrawal,
            'fields': ['withdrawal_id', 'member__name', 'amount', 'purpose', 'status', 'application_date', 'approval_date', 'payment_method'],
            'headers': ['Withdrawal ID', 'Member', 'Amount', 'Purpose', 'Status', 'Application Date', 'Approval Date', 'Payment Method'],
            'title': 'Withdrawals Report'
        },
        'payments': {
            'model': Payment,
            'fields': ['payment_id', 'loan__member__name', 'loan__loan_id', 'amount', 'payment_date', 'payment_method'],
            'headers': ['Payment ID', 'Member', 'Loan ID', 'Amount', 'Payment Date', 'Payment Method'],
            'title': 'Payments Report'
        }
    }

    if model_name not in model_configs:
        messages.error(request, 'Invalid export type')
        return redirect('sacco:dashboard')

    config = model_configs[model_name]
    queryset = config['model'].objects.all().order_by('-pk')

    if hasattr(config['model'], 'member'):
        queryset = queryset.select_related('member')
    if model_name == 'payments':
        queryset = queryset.select_related('loan__member')

    data = []
    for obj in queryset:
        row = []
        for field in config['fields']:
            value = obj
            for attr in field.split('__'):
                value = getattr(value, attr, '')
            if hasattr(value, 'isoformat') and value:  # Date/Time fields
                value = value.isoformat()
            row.append(str(value) if value else '')
        data.append(row)

    filename = f"{model_name}_report_{timezone.now().date()}"

    if format_type == 'csv':
        return export_csv(data, config['headers'], filename)
    elif format_type == 'pdf':
        try:
            return export_pdf(data, config['headers'], config['title'], filename)
        except ImportError as e:
            messages.error(request, f"PDF export not available: {str(e)}")
            return redirect('sacco:dashboard')
    elif format_type == 'excel':
        try:
            return export_excel(data, config['headers'], config['title'], filename)
        except ImportError as e:
            messages.error(request, f"Excel export not available: {str(e)}")
            return redirect('sacco:dashboard')
    elif format_type == 'docx':
        try:
            return export_docx(data, config['headers'], config['title'], filename)
        except ImportError as e:
            messages.error(request, f"DOCX export not available: {str(e)}")
            return redirect('sacco:dashboard')
    else:
        messages.error(request, 'Invalid export format')
        return redirect('sacco:dashboard')

def export_csv(data, headers, filename):
    """Export data as CSV"""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'

    writer = csv.writer(response)
    writer.writerow(headers)
    writer.writerows(data)
    return response

def export_pdf(data, headers, title, filename):
    """Export data as PDF"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab is required for PDF export. Install with: pip install reportlab")

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    styles = getSampleStyleSheet()
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Prepare table data
    table_data = [headers] + data

    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))

    elements.append(table)
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    return response

def export_excel(data, headers, title, filename):
    """Export data as Excel"""
    if not OPENPYXL_AVAILABLE:
        raise ImportError("openpyxl is required for Excel export. Install with: pip install openpyxl")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = title

    # Add title
    ws['A1'] = title
    ws['A1'].font = ws['A1'].font.copy(bold=True, size=16)

    # Add headers
    for col_num, header in enumerate(headers, 1):
        ws.cell(row=3, column=col_num, value=header)
        ws.cell(row=3, column=col_num).font = ws.cell(row=3, column=col_num).font.copy(bold=True)

    # Add data
    for row_num, row_data in enumerate(data, 4):
        for col_num, cell_value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=cell_value)

    wb.save(response)
    return response

def export_docx(data, headers, title, filename):
    """Export data as DOCX"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'

    document = Document()
    document.add_heading(title, 0)

    # Create table
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Add data
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)

    document.save(response)
    return response

def login_view(request):
    """Login view for members using phone number as username"""
    if request.user.is_authenticated:
        return redirect('sacco:dashboard')

    if request.method == 'POST':
        phone = request.POST.get('identifier')  # phone number as username
        password = request.POST.get('password')

        user = authenticate(request, username=phone, password=password)

        if user is not None:
            login(request, user)
            messages.success(request, f'Welcome back, {user.get_full_name()}!')
            return redirect('sacco:dashboard')
        else:
            messages.error(request, 'Invalid phone number or password. Please try again.')

    return render(request, 'sacco/login.html')

def forgot_password(request):
    """Forgot password view - send OTP to phone number"""
    if request.user.is_authenticated:
        return redirect('sacco:dashboard')

    if request.method == 'POST':
        form = ForgotPasswordForm(request.POST)
        if form.is_valid():
            phone = form.cleaned_data['phone']

            # Check if member exists with this phone number
            try:
                member = Member.objects.get(phone=phone)
                # Generate OTP
                otp_code = ''.join(random.choices(string.digits, k=6))
                expires_at = timezone.now() + timedelta(minutes=10)

                # Save OTP
                OTP.objects.create(
                    phone=phone,
                    otp_code=otp_code,
                    expires_at=expires_at,
                    purpose='password_reset'
                )

                # Send SMS (for now, just print to console - in production, integrate with SMS service)
                print(f"OTP for {phone}: {otp_code}")  # Replace with actual SMS sending
                send_sms(phone, f"Your SACCO OTP code is: {otp_code}. Valid for 10 minutes.")

                messages.success(request, 'OTP sent to your phone number. Please check your messages.')
                return redirect('sacco:verify_otp', phone=phone)

            except Member.DoesNotExist:
                messages.error(request, 'No account found with this phone number.')

    else:
        form = ForgotPasswordForm()

    return render(request, 'sacco/forgot_password.html', {'form': form})

def verify_otp(request, phone):
    """Verify OTP and allow password reset"""
    if request.user.is_authenticated:
        return redirect('sacco:dashboard')

    if request.method == 'POST':
        form = OTPVerificationForm(request.POST)
        if form.is_valid():
            otp_code = form.cleaned_data['otp_code']

            # Check if OTP exists and is valid
            try:
                otp = OTP.objects.get(
                    phone=phone,
                    otp_code=otp_code,
                    is_used=False,
                    purpose='password_reset'
                )

                if otp.is_expired():
                    messages.error(request, 'OTP has expired. Please request a new one.')
                    return redirect('sacco:forgot_password')

                # Mark OTP as used
                otp.is_used = True
                otp.save()

                # Redirect to set password
                messages.success(request, 'OTP verified successfully. Please set your new password.')
                return redirect('sacco:set_password', phone=phone)

            except OTP.DoesNotExist:
                messages.error(request, 'Invalid OTP code.')

    else:
        form = OTPVerificationForm()

    return render(request, 'sacco/verify_otp.html', {'form': form, 'phone': phone})

def set_password(request, phone):
    """Set new password for the user"""
    if request.user.is_authenticated:
        return redirect('sacco:dashboard')

    try:
        member = Member.objects.get(phone=phone)
    except Member.DoesNotExist:
        messages.error(request, 'Account not found.')
        return redirect('sacco:forgot_password')

    if request.method == 'POST':
        form = SetPasswordForm(request.POST)
        if form.is_valid():
            password = form.cleaned_data['password']

            # Create or update user account
            if member.user:
                user = member.user
            else:
                # Create new user account
                user = User.objects.create_user(
                    username=phone,
                    email=member.email or '',
                    first_name=member.name.split()[0] if member.name else '',
                    last_name=' '.join(member.name.split()[1:]) if member.name and len(member.name.split()) > 1 else ''
                )
                member.user = user
                member.save()

            # Set new password
            user.set_password(password)
            user.save()

            messages.success(request, 'Password set successfully. You can now login.')
            return redirect('sacco:login')

    else:
        form = SetPasswordForm()

    return render(request, 'sacco/set_password.html', {'form': form, 'phone': phone})

def logout_view(request):
    """Logout view"""
    logout(request)
    messages.info(request, 'You have been logged out.')
    return redirect('sacco:login')

def dashboard(request):
    """Dashboard view - shows role-specific content"""
    if not request.user.is_authenticated:
        return redirect('sacco:login')

    user_role = get_user_role(request.user)

    # Common statistics for all roles
    total_members = Member.objects.count()
    deposits = Savings.objects.filter(transaction_type='deposit').aggregate(Sum('amount'))['amount__sum'] or 0
    withdrawals = Savings.objects.filter(transaction_type='withdrawal').aggregate(Sum('amount'))['amount__sum'] or 0
    total_savings = deposits - withdrawals
    # Member activity counts: consider members with linked User and active flag as active
    active_members = Member.objects.filter(user__is_active=True).count()
    inactive_members = total_members - active_members
    total_loans_given_out = Loan.objects.filter(status='approved').aggregate(Sum('amount'))['amount__sum'] or 0
    pending_loans = Loan.objects.filter(status='pending').count()
    pending_loans_amount = Loan.objects.filter(status='pending').aggregate(Sum('amount'))['amount__sum'] or 0
    approved_loans = Loan.objects.filter(status='approved').count()
    rejected_loans = Loan.objects.filter(status='rejected').count()
    total_loans = Loan.objects.count()
    # Withdrawals summary
    total_withdrawals_amount = Withdrawal.objects.aggregate(total=Sum('amount'))['total'] or 0
    pending_withdrawals_count = Withdrawal.objects.filter(status='pending').count()
    approved_withdrawals_count = Withdrawal.objects.filter(status='approved').count()
    rejected_withdrawals_count = Withdrawal.objects.filter(status='rejected').count()
    total_withdrawals_count = Withdrawal.objects.count()

    # Calculate overdue loans
    overdue_loans = 0
    today = timezone.now().date()
    for loan in Loan.objects.filter(status='approved'):
        if loan.approval_date:
            due_date = loan.approval_date + timedelta(days=loan.term_months * 30)
            if due_date < today:
                overdue_loans += 1

    # Role-specific content
    if user_role == 'user':
        # User dashboard - personal information and their own data
        try:
            member = Member.objects.get(user=request.user)
            deposits = Savings.objects.filter(member=member, transaction_type='deposit').aggregate(Sum('amount'))['amount__sum'] or Decimal('0')
            withdrawals = Savings.objects.filter(member=member, transaction_type='withdrawal').aggregate(Sum('amount'))['amount__sum'] or Decimal('0')
            user_savings = float(deposits - withdrawals)

            user_loans = Loan.objects.filter(member=member)
            approved_loans_total = sum(float(loan.amount) for loan in user_loans.filter(status='approved'))
            payments_total = sum(float(payment.amount) for payment in Payment.objects.filter(loan__member=member))
            user_loan_balance = approved_loans_total - payments_total

            user_pending_loans = user_loans.filter(status='pending').count()

            # Recent activities for user
            recent_savings = list(Savings.objects.filter(member=member).order_by('-transaction_date')[:3])
            recent_loans = list(Loan.objects.filter(member=member).order_by('-application_date')[:3])
            recent_activities = recent_savings + recent_loans

            # Sort activities safely
            def get_activity_date(activity):
                if hasattr(activity, 'transaction_date') and activity.transaction_date:
                    return activity.transaction_date
                elif hasattr(activity, 'application_date') and activity.application_date:
                    return activity.application_date
                else:
                    return today

            recent_activities.sort(key=get_activity_date, reverse=True)
            recent_activities = recent_activities[:5]

            # Get all savings transactions for the member
            member_recent_savings = Savings.objects.filter(member=member).order_by('-transaction_date')
            
            # Get total savings made by member (from admin contributions/deposits)
            admin_deposits = Savings.objects.filter(member=member, transaction_type='deposit').aggregate(Sum('amount'))['amount__sum'] or Decimal('0')
            admin_withdrawals = Savings.objects.filter(member=member, transaction_type='withdrawal').aggregate(Sum('amount'))['amount__sum'] or Decimal('0')

            context = {
                'user_role': 'user',
                'member': member,
                'member_savings': user_savings,
                'member_loans': user_loans,
                'member_recent_savings': member_recent_savings,
                'admin_deposits': admin_deposits,
                'admin_withdrawals': admin_withdrawals,
                'user_loan_balance': user_loan_balance,
                'user_pending_loans': user_pending_loans,
                'recent_activities': recent_activities,

                # Member presence statistics (for dashboards/widgets)
                'total_members': total_members,
                'active_members': active_members,
                'inactive_members': inactive_members,
            }
        except Member.DoesNotExist:
            context = {'user_role': 'user', 'error': 'Member profile not found', 'recent_activities': []}
        except Exception as e:
            # Catch any other unexpected errors
            context = {'user_role': 'user', 'error': f'An error occurred: {str(e)}', 'recent_activities': []}

    elif user_role in ['admin', 'superadmin']:
        # Admin/Superadmin dashboard - full statistics and management overview
        recent_savings = list(Savings.objects.select_related('member').order_by('-transaction_date')[:3])
        recent_loans = list(Loan.objects.select_related('member').order_by('-application_date')[:3])
        recent_activities = recent_savings + recent_loans
        recent_activities.sort(key=lambda x: getattr(x, 'transaction_date', getattr(x, 'application_date', today)), reverse=True)
        recent_activities = recent_activities[:5]

        # Superadmin-only: pending notifications count for dashboard widget
        pending_count = 0
        if user_role == 'superadmin':
            pending_count = Notification.objects.filter(status='pending').count()

        context = {
            'user_role': user_role,
            'is_admin': True,
            'total_members': total_members,
            'active_members': active_members,
            'inactive_members': inactive_members,
            'total_loans': total_loans,
            'total_savings': total_savings,
            'deposits_total': deposits,
            'withdrawals_total': withdrawals,
            'total_withdrawals_amount': total_withdrawals_amount,
            'total_withdrawals_count': total_withdrawals_count,
            'pending_withdrawals_count': pending_withdrawals_count,
            'approved_withdrawals_count': approved_withdrawals_count,
            'rejected_withdrawals_count': rejected_withdrawals_count,
            'total_loans_given_out': total_loans_given_out,
            'pending_loans': pending_loans,
            'pending_loans_amount': pending_loans_amount,
            'approved_loans': approved_loans,
            'rejected_loans': rejected_loans,
            'overdue_loans': overdue_loans,
            'pending_count': pending_count,
            'recent_savings': recent_savings,
            'recent_loans': recent_loans,
            'recent_activities': recent_activities,
        }

    else:
        # Default fallback
        context = {'error': 'Role not recognized'}

    # Render appropriate template based on user role
    if user_role == 'user':
        return render(request, 'sacco/member_dashboard.html', context)
    else:
        return render(request, 'sacco/dashboard.html', context)

@login_required
def settings(request):
    """Settings view - shows user profile and admin options"""
    try:
        member = Member.objects.get(user=request.user)
        member_exists = True
    except Member.DoesNotExist:
        member = None
        member_exists = False
    is_admin = request.user.is_staff or request.user.is_superuser

    if request.method == 'POST':
        if 'update_profile' in request.POST and member_exists:
            # Update member profile
            form = MemberForm(request.POST, instance=member)
            if form.is_valid():
                form.save()
                messages.success(request, 'Profile updated successfully!')
                return redirect('sacco:settings')
        elif 'change_password' in request.POST:
            # Change password
            current_password = request.POST.get('current_password')
            new_password = request.POST.get('new_password')
            confirm_password = request.POST.get('confirm_password')

            if not request.user.check_password(current_password):
                messages.error(request, 'Current password is incorrect.')
            elif new_password != confirm_password:
                messages.error(request, 'New passwords do not match.')
            elif len(new_password) < 8:
                messages.error(request, 'Password must be at least 8 characters long.')
            else:
                request.user.set_password(new_password)
                request.user.save()
                messages.success(request, 'Password changed successfully! Please login again.')
                return redirect('sacco:login')
        elif 'update_user_settings' in request.POST and is_admin:
            # Admin settings
            pass  # Add admin-specific settings here

    return render(request, 'sacco/settings.html', {
        'member': member,
        'member_exists': member_exists,
        'is_admin': is_admin
    })

def csrf_failure(request, reason=""):
    """Custom CSRF failure view"""
    logger.warning(f"CSRF failure: {reason}, User: {request.user}, Path: {request.path}")
    return render(request, 'sacco/csrf_failure.html', {'reason': reason})


@login_required
def loan_calculator(request):
    """Render loan calculator page (frontend calls api_loan_calculator)."""
    return render(request, 'sacco/loan_calculator.html')


@user_required
def apply_loan(request):
    """Apply for a loan by member - displays loan details and progress first"""
    try:
        member = Member.objects.get(user=request.user)
    except Member.DoesNotExist:
        messages.error(request, 'You are not a registered member.')
        return redirect('sacco:dashboard')

    # Get member's loan details
    member_loans = Loan.objects.filter(member=member).order_by('-application_date')
    approved_loans_amount = sum(float(loan.amount) for loan in member_loans.filter(status='approved'))
    payments_total = sum(float(payment.amount) for payment in Payment.objects.filter(loan__member=member))
    loan_balance = approved_loans_amount - payments_total
    pending_loans = member_loans.filter(status='pending').count()

    # Get payment progress for each approved loan
    approved_loans = member_loans.filter(status='approved')
    loan_progress_details = []
    
    for loan in approved_loans:
        loan_payments = Payment.objects.filter(loan=loan).aggregate(total=Sum('amount'))['total'] or 0
        principal_amount = float(loan.amount)
        amount_paid = float(loan_payments)
        remaining_balance = principal_amount - amount_paid
        payment_percentage = (amount_paid / principal_amount * 100) if principal_amount > 0 else 0
        
        loan_progress_details.append({
            'loan': loan,
            'principal_amount': principal_amount,
            'amount_paid': amount_paid,
            'remaining_balance': remaining_balance,
            'payment_percentage': round(payment_percentage, 1),
            'total_payments': Payment.objects.filter(loan=loan).count(),
        })

    if request.method == 'POST':
        form = LoanApplicationForm(request.POST)
        if form.is_valid():
            try:
                loan = form.save(commit=False)
                loan.member = member
                loan.save()
                
                # Create notification for admin
                Notification.objects.create(
                    member=member,
                    notification_type='loan_application',
                    title=f'Loan Application from {member.name}',
                    message=f'New loan application: Amount: UGX {loan.amount}, Term: {loan.term_months} months, Purpose: {loan.purpose}',
                    loan=loan,
                )
                
                messages.success(request, 'Loan application submitted successfully!')
                logger.info(f"Loan application created for member {member.name}, Amount: {loan.amount}")
                return redirect('sacco:dashboard')
            except Exception as e:
                logger.error(f"Error creating loan application: {str(e)}", exc_info=True)
                messages.error(request, f'An error occurred: {str(e)}')
        else:
            logger.warning(f"Loan form validation failed: {form.errors}")
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {str(error)}')
    else:
        form = LoanApplicationForm()

    context = {
        'form': form,
        'member_loans': member_loans,
        'loan_balance': loan_balance,
        'pending_loans': pending_loans,
        'approved_loans_amount': approved_loans_amount,
        'loan_progress_details': loan_progress_details,
        'has_approved_loans': len(loan_progress_details) > 0,
    }
    return render(request, 'sacco/apply_loan.html', context)

@user_required
def make_payment(request):
    """Make a payment for a loan - displays payment details first"""
    member = get_object_or_404(Member, user=request.user)

    # Get member's payment and savings details
    member_loans = Loan.objects.filter(member=member, status='approved').order_by('-application_date')
    payment_history = Payment.objects.filter(loan__member=member).select_related('loan').order_by('-payment_date')[:10]
    
    # Get savings details with Decimal safety
    deposits = Savings.objects.filter(member=member, transaction_type='deposit').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    withdrawals = Savings.objects.filter(member=member, transaction_type='withdrawal').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    current_savings = deposits - withdrawals
    
    approved_loans_qs = Loan.objects.filter(member=member, status='approved')
    has_approved_loans = approved_loans_qs.exists()
    
    if request.method == 'POST':
        form = PaymentForm(request.POST)
        form.fields['loan'].queryset = approved_loans_qs  # Restrict to user's loans
        if form.is_valid():
            try:
                payment = form.save(commit=False)
                if not payment.loan:
                    raise ValueError("Please select a loan")
                if not payment.payment_date:
                    payment.payment_date = timezone.now().date()
                payment.save()
                logger.info(f"Payment created successfully for member {member.name}, Loan: {payment.loan.loan_id}, Amount: {payment.amount}")
                messages.success(request, 'Payment made successfully!')
                return redirect('sacco:dashboard')
            except IntegrityError as ie:
                logger.error(f"DB error saving payment: {str(ie)}")
                messages.error(request, "Payment could not be saved. Please check loan selection.")
            except ValueError as ve:
                logger.error(f"Value error in payment: {str(ve)}")
                messages.error(request, str(ve))
            except Exception as e:
                logger.error(f"Error saving payment for member {member.name}: {str(e)}", exc_info=True)
                messages.error(request, f'An error occurred while processing payment: {str(e)}')
        else:
            logger.warning(f"Payment form validation failed for member {member.name}: {form.errors}")
            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, f'{field}: {str(error)}')
    else:
        if has_approved_loans:
            form = PaymentForm()
            form.fields['loan'].queryset = approved_loans_qs
            form.fields['payment_date'].initial = timezone.now().date()
        else:
            form = None
            messages.info(request, "No approved loans available for payment.")

    context = {
        'form': form,
        'member_loans': member_loans,
        'payment_history': payment_history,
        'current_savings': current_savings,
        'deposits': deposits,
        'withdrawals': withdrawals,
        'has_approved_loans': has_approved_loans,
    }
    return render(request, 'sacco/make_payment.html', context)

@user_required
def apply_withdrawal(request):
    """Apply for a withdrawal by member - displays withdrawal details first"""
    member = get_object_or_404(Member, user=request.user)

    # Calculate current balance with Decimal safety
    deposits = Savings.objects.filter(member=member, transaction_type='deposit').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    withdrawals_total = Savings.objects.filter(member=member, transaction_type='withdrawal').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    current_balance = deposits - withdrawals_total
    
    # Get withdrawal history - filter BEFORE slicing
    all_withdrawals = Withdrawal.objects.filter(member=member).order_by('-application_date')
    approved_withdrawals = all_withdrawals.filter(status='approved').count()
    pending_withdrawals = all_withdrawals.filter(status='pending').count()
    withdrawal_history = all_withdrawals[:10]  # Slice after filtering for counts
    
    form = None

    if request.method == 'POST':
        form = WithdrawalApplicationForm(request.POST)
        if form.is_valid():
            try:
                withdrawal = form.save(commit=False)
                withdrawal_amount = withdrawal.amount
                withdrawal.member = member
                withdrawal.status = 'pending'
                if withdrawal.application_date is None:
                    withdrawal.application_date = timezone.now().date()
                    
                if withdrawal_amount > current_balance:
                    messages.error(request, f'Insufficient balance. Your current balance is UGX {float(current_balance):.0f}.')
                else:
                    withdrawal.save()
                    
                    # Create notification safely
                    try:
                        Notification.objects.create(
                            member=member,
                            notification_type='withdrawal_application',
                            title=f'Withdrawal Application from {member.name}',
                            message=f'New withdrawal application: Amount: UGX {withdrawal_amount}, Purpose: {withdrawal.purpose}',
                            withdrawal=withdrawal,
                        )
                    except Exception as notif_error:
                        logger.error(f"Error creating withdrawal notification: {str(notif_error)}")
                    
                    messages.success(request, f'Withdrawal application submitted successfully! Amount: UGX {float(withdrawal_amount):.0f}')
                    return redirect('sacco:dashboard')
            except IntegrityError as ie:
                logger.error(f"DB error saving withdrawal: {str(ie)}")
                messages.error(request, "Withdrawal could not be saved. Please check amount.")
            except Exception as e:
                logger.error(f"Error processing withdrawal: {str(e)}", exc_info=True)
                messages.error(request, f'An error occurred: {str(e)}')
        else:
            logger.warning(f"Withdrawal form errors: {form.errors}")
            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, f'{field}: {str(error)}')
    
    # Initialize form if not already done
    if form is None:
        form = WithdrawalApplicationForm()

    context = {
        'form': form,
        'current_balance': current_balance,
        'deposits': deposits,
        'withdrawals_total': withdrawals_total,
        'withdrawal_history': withdrawal_history,
        'approved_withdrawals': approved_withdrawals,
        'pending_withdrawals': pending_withdrawals,
    }
    return render(request, 'sacco/apply_withdrawal.html', context)

@login_required
def withdrawals_list(request):
    """List all withdrawals (admin view)"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')
    withdrawals = Withdrawal.objects.all().select_related('member').order_by('-application_date')
    return render(request, 'sacco/withdrawals_list.html', {'withdrawals': withdrawals})


# ==================== NOTIFICATION VIEWS ====================

@login_required
def notifications_list(request):
    """View all notifications for the current user"""
    try:
        member = Member.objects.get(user=request.user)
        is_admin = request.user.is_staff or request.user.is_superuser
    except Member.DoesNotExist:
        is_admin = request.user.is_staff or request.user.is_superuser
        member = None

    if is_admin:
        # Admin: show all pending notifications sorted by date
        notifications = Notification.objects.all().select_related('member', 'loan', 'withdrawal').order_by('-created_at')
        unread_count = notifications.filter(status='pending').count()
    else:
        # Member: show only their notifications
        notifications = Notification.objects.filter(member=member).select_related('loan', 'withdrawal').order_by('-created_at')
        unread_count = notifications.filter(status='pending').count()

    # Do NOT auto-mark as read here - let user mark as read when viewing the notification

    context = {
        'notifications': notifications,
        'unread_count': unread_count,
        'is_admin': is_admin,
    }
    return render(request, 'sacco/notifications_list.html', context)


@login_required
def notification_detail(request, notification_id):
    """View a single notification and optionally add a reply (admin only)"""
    notification = get_object_or_404(Notification, pk=notification_id)
    user = request.user
    is_admin = user.is_staff or user.is_superuser

    # Check permissions
    try:
        member = Member.objects.get(user=user)
    except Member.DoesNotExist:
        member = None

    # Only allow member to view their own notifications, or admins to view all
    if not is_admin and member != notification.member:
        messages.error(request, 'You do not have permission to view this notification.')
        return redirect('sacco:notifications_list')

    # Mark as read
    notification.status = 'read'
    notification.save()

    # Handle reply submission (admin only)
    if request.method == 'POST' and is_admin:
        reply_message = request.POST.get('reply_message')
        is_internal = request.POST.get('is_internal') == 'on'
        mark_resolved = request.POST.get('mark_resolved') == 'on'

        if reply_message:
            NotificationReply.objects.create(
                notification=notification,
                admin_user=user,
                reply_message=reply_message,
                is_internal=is_internal,
            )

            if mark_resolved:
                notification.status = 'resolved'
                notification.admin_user = user
                notification.save()
                messages.success(request, 'Reply sent and notification marked as resolved.')
            else:
                messages.success(request, 'Reply sent successfully.')

            return redirect('sacco:notification_detail', notification_id=notification_id)

    # Get all replies
    replies = notification.replies.all().select_related('admin_user')

    context = {
        'notification': notification,
        'replies': replies,
        'is_admin': is_admin,
    }
    return render(request, 'sacco/notification_detail.html', context)


@admin_required
def admin_notifications(request):
    """Admin dashboard for managing notifications (admins only)"""
    # Superadmins should use the superadmin_notifications view instead
    if request.user.is_superuser:
        return redirect('sacco:superadmin_notifications')
    
    # Get all notifications sorted by status and date (for regular admins)
    pending_notifications = Notification.objects.filter(status='pending').select_related('member', 'loan', 'withdrawal').order_by('-created_at')
    read_notifications = Notification.objects.filter(status='read').select_related('member', 'loan', 'withdrawal').order_by('-created_at')
    resolved_notifications = Notification.objects.filter(status='resolved').select_related('member', 'loan', 'withdrawal').order_by('-created_at')

    context = {
        'pending_notifications': pending_notifications,
        'read_notifications': read_notifications,
        'resolved_notifications': resolved_notifications,
        'pending_count': pending_notifications.count(),
        'read_count': read_notifications.count(),
    }
    return render(request, 'sacco/admin_notifications.html', context)


@login_required
def superadmin_notifications(request):
    """Superadmin dashboard for managing ALL notifications from members and admins"""
    user = request.user
    if not user.is_superuser:
        messages.error(request, 'Permission denied. Only superadmins can access this.')
        return redirect('sacco:admin_notifications' if user.is_staff else 'sacco:dashboard')
    
    # Get all notifications by status and date
    pending_notifications = Notification.objects.filter(status='pending').select_related('member', 'loan', 'withdrawal', 'admin_user').order_by('-created_at')
    read_notifications = Notification.objects.filter(status='read').select_related('member', 'loan', 'withdrawal', 'admin_user').order_by('-created_at')
    resolved_notifications = Notification.objects.filter(status='resolved').select_related('member', 'loan', 'withdrawal', 'admin_user').order_by('-created_at')
    
    # Get statistics
    total_notifications = Notification.objects.count()
    total_pending = pending_notifications.count()
    total_read = read_notifications.count()
    total_resolved = resolved_notifications.count()
    
    # Get notification type breakdown
    admin_messages = Notification.objects.filter(notification_type='admin_message').count()
    loan_related = Notification.objects.filter(notification_type__in=['loan_application', 'loan_approval', 'loan_rejection']).count()
    withdrawal_related = Notification.objects.filter(notification_type__in=['withdrawal_application', 'withdrawal_approval', 'withdrawal_rejection']).count()
    
    context = {
        'pending_notifications': pending_notifications,
        'read_notifications': read_notifications,
        'resolved_notifications': resolved_notifications,
        'pending_count': total_pending,
        'read_count': total_read,
        'resolved_count': total_resolved,
        'total_notifications': total_notifications,
        'admin_messages_count': admin_messages,
        'loan_related_count': loan_related,
        'withdrawal_related_count': withdrawal_related,
        'is_superadmin': True,
    }
    return render(request, 'sacco/superadmin_notifications.html', context)


@login_required
def create_notification(request):
    """Create and send a notification to a member or all members (admin/superadmin only)"""
    user = request.user
    if not (user.is_staff or user.is_superuser):
        messages.error(request, 'Permission denied.')
        return redirect('sacco:dashboard')

    if request.method == 'POST':
        form = NotificationForm(request.POST)
        if form.is_valid():
            try:
                send_to_all = form.cleaned_data.get('send_to_all_members')
                selected_member = form.cleaned_data.get('member')
                title = form.cleaned_data.get('title')
                message = form.cleaned_data.get('message')

                if send_to_all:
                    # Get all members
                    all_members = Member.objects.all()
                    notification_count = 0

                    # Create notification for each member
                    for member in all_members:
                        Notification.objects.create(
                            member=member,
                            notification_type='admin_message',
                            title=title,
                            message=message,
                            status='pending',
                            admin_user=user
                        )
                        notification_count += 1

                    logger.info(f"Broadcast notification created by {user.username} to {notification_count} members: {title}")
                    messages.success(request, f"Notification sent to all {notification_count} members successfully.")
                else:
                    # Send to single member
                    Notification.objects.create(
                        member=selected_member,
                        notification_type='admin_message',
                        title=title,
                        message=message,
                        status='pending',
                        admin_user=user
                    )

                    logger.info(f"Notification created by {user.username} for member {selected_member.name}: {title}")
                    messages.success(request, f"Notification sent to {selected_member.name} successfully.")

                return redirect('sacco:admin_notifications')
            except Exception as e:
                logger.error(f"Error creating notification: {str(e)}")
                messages.error(request, f"Error creating notification: {str(e)}")
    else:
        form = NotificationForm()

    context = {
        'form': form,
        'page_title': 'Send Notification to Member',
    }
    return render(request, 'sacco/create_notification.html', context)


@login_required
def mark_notification_resolved(request, notification_id):
    """Mark a notification as resolved (admin only)"""
    user = request.user
    if not (user.is_staff or user.is_superuser):
        messages.error(request, 'Permission denied.')
        return redirect('sacco:dashboard')

    notification = get_object_or_404(Notification, pk=notification_id)
    notification.status = 'resolved'
    notification.admin_user = user
    notification.save()

    messages.success(request, 'Notification marked as resolved.')
    return redirect('sacco:notification_detail', notification_id=notification_id)

@login_required
def add_withdrawal(request):
    """Add a new withdrawal (admin only)"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')

    if request.method == 'POST':
        form = WithdrawalForm(request.POST)
        if form.is_valid():
            try:
                withdrawal = form.save()
                logger.info(f"Withdrawal added by admin: {withdrawal.withdrawal_id} for member {withdrawal.member.name}, Amount: {withdrawal.amount}, Status: {withdrawal.status}")
                messages.success(request, 'Withdrawal added successfully!')
                return redirect('sacco:withdrawals_list')
            except Exception as e:
                logger.error(f"Error saving withdrawal: {str(e)}", exc_info=True)
                messages.error(request, f'An error occurred while saving the withdrawal: {str(e)}')
        else:
            logger.warning(f"Withdrawal form validation failed: {form.errors}")
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {str(error)}')
    else:
        form = WithdrawalForm()

    return render(request, 'sacco/add_withdrawal.html', {'form': form})

@login_required
def edit_withdrawal(request, withdrawal_id):
    """Edit an existing withdrawal"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')

    withdrawal = get_object_or_404(Withdrawal, pk=withdrawal_id)

    if request.method == 'POST':
        form = WithdrawalForm(request.POST, instance=withdrawal)
        if form.is_valid():
            updated_withdrawal = form.save(commit=False)
            if withdrawal.status != updated_withdrawal.status:
                if updated_withdrawal.status == 'approved':
                    # Check if member has sufficient balance
                    deposits = Savings.objects.filter(
                        member=withdrawal.member,
                        transaction_type='deposit'
                    ).aggregate(Sum('amount'))['amount__sum'] or 0

                    withdrawals = Savings.objects.filter(
                        member=withdrawal.member,
                        transaction_type='withdrawal'
                    ).aggregate(Sum('amount'))['amount__sum'] or 0

                    current_balance = deposits - withdrawals

                    if current_balance >= withdrawal.amount:
                        updated_withdrawal.approval_date = timezone.now().date()
                        updated_withdrawal.approved_by = request.user

                        # Create a savings transaction for the withdrawal
                        Savings.objects.create(
                            member=withdrawal.member,
                            amount=withdrawal.amount,
                            transaction_type='withdrawal',
                            payment_method=withdrawal.payment_method,
                            transaction_date=timezone.now().date()
                        )

                        messages.success(request, f'Withdrawal approved successfully! Amount: UGX {withdrawal.amount}')
                        # Send notification to member (TODO: implement send_withdrawal_notification if needed)
                    else:
                        messages.error(request, f'Insufficient balance. Current balance: UGX {current_balance}, Requested: UGX {withdrawal.amount}')
                        return redirect('sacco:edit_withdrawal', withdrawal_id=withdrawal_id)

                elif updated_withdrawal.status == 'rejected':
                    updated_withdrawal.approval_date = timezone.now().date()
                    updated_withdrawal.approved_by = request.user
                    messages.success(request, 'Withdrawal rejected.')
                    # Send notification to member (TODO: implement send_withdrawal_notification if needed)
                else:
                    updated_withdrawal.approval_date = None
                    updated_withdrawal.approved_by = None
            updated_withdrawal.save()
            return redirect('sacco:withdrawals_list')
    else:
        form = WithdrawalForm(instance=withdrawal)

    return render(request, 'sacco/edit_withdrawal.html', {'form': form, 'withdrawal': withdrawal})

@login_required
def delete_withdrawal(request, withdrawal_id):
    """Delete a withdrawal"""
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Permission denied.")
        return redirect('sacco:dashboard')

    withdrawal = get_object_or_404(Withdrawal, pk=withdrawal_id)

    if request.method == 'POST':
        withdrawal.delete()
        messages.success(request, 'Withdrawal deleted successfully!')
        return redirect('sacco:withdrawals_list')

    return render(request, 'sacco/delete_withdrawal_confirm.html', {'withdrawal': withdrawal})

def contact(request):
    """Contact information for sacco staff"""
    return render(request, 'sacco/contact.html')

# add this function so URL reversing for 'settings' works
def settings_view(request):
    """
    Simple settings view placeholder.
    Change to render an actual settings template if needed.
    """
    return redirect('sacco:dashboard')


# ==================== API ENDPOINTS ====================

@login_required
@csrf_exempt
def api_loan_calculator(request):
    """
    API endpoint for loan calculator
    POST parameters: loan_amount, interest_rate, loan_term_months
    Returns: JSON with monthly payment, total interest, total repayment
    """
    if request.method == 'POST':
        try:
            # IMPORTANT: some browsers/clients may send JSON as bytes; request.body is safe.
            data = json.loads(request.body.decode('utf-8') if isinstance(request.body, (bytes, bytearray)) else request.body)

            loan_amount = float(data.get('loan_amount', 0))
            interest_rate = float(data.get('interest_rate', 0))
            loan_term_months = int(data.get('loan_term_months', 0))

            # Validate inputs
            if loan_amount <= 0 or loan_term_months <= 0:
                return JsonResponse({
                    'success': False,
                    'error': 'Loan amount and term must be greater than 0'
                }, status=400)

            # Monthly interest rate
            monthly_interest_rate = interest_rate / 100 / 12

            # Amortization / EMI formula
            if monthly_interest_rate > 0:
                compound = (1 + monthly_interest_rate) ** loan_term_months
                monthly_payment = (loan_amount * monthly_interest_rate * compound) / (compound - 1)
            else:
                monthly_payment = loan_amount / loan_term_months

            total_repayment = monthly_payment * loan_term_months
            total_interest = total_repayment - loan_amount

            return JsonResponse({
                'success': True,
                'monthly_payment': round(monthly_payment, 2),
                'total_interest': round(total_interest, 2),
                'total_repayment': round(total_repayment, 2),
                'loan_amount': loan_amount,
                'interest_rate': interest_rate,
                'loan_term_months': loan_term_months
            }, status=200)

        except json.JSONDecodeError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON format'
            }, status=400)
        except Exception as e:
            # Include exception type for easier debugging
            return JsonResponse({
                'success': False,
                'error': f'{type(e).__name__}: {str(e)}'
            }, status=400)

    return JsonResponse({
        'success': False,
        'error': 'Only POST requests are allowed'
    }, status=405)



@login_required
def api_withdrawal_apply(request):
    """
    API endpoint for withdrawal application
    POST parameters: amount, purpose, payment_method
    Returns: JSON with status and withdrawal details
    """
    if request.method == 'POST':
        try:
            member = Member.objects.get(user=request.user)
        except Member.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'You are not a registered member'
            }, status=403)
        
        try:
            data = json.loads(request.body)
            
            amount = Decimal(str(data.get('amount', 0)))
            purpose = data.get('purpose', '').strip()
            payment_method = data.get('payment_method', 'cash')
            
            if amount <= 0:
                return JsonResponse({
                    'success': False,
                    'error': 'Amount must be greater than 0'
                }, status=400)
            
            if not purpose:
                return JsonResponse({
                    'success': False,
                    'error': 'Purpose is required'
                }, status=400)
            
            # Calculate current balance
            deposits = Savings.objects.filter(
                member=member, 
                transaction_type='deposit'
            ).aggregate(Sum('amount'))['amount__sum'] or 0
            
            withdrawals = Savings.objects.filter(
                member=member, 
                transaction_type='withdrawal'
            ).aggregate(Sum('amount'))['amount__sum'] or 0
            
            current_balance = deposits - withdrawals
            
            if amount > current_balance:
                return JsonResponse({
                    'success': False,
                    'error': f'Insufficient balance. Current balance: UGX {float(current_balance):.2f}'
                }, status=400)
            
            # Create withdrawal
            withdrawal = Withdrawal.objects.create(
                member=member,
                amount=amount,
                purpose=purpose,
                payment_method=payment_method,
                status='pending'
            )
            
            # Create notification
            Notification.objects.create(
                member=member,
                notification_type='withdrawal_application',
                title=f'Withdrawal Application Submitted',
                message=f'Your withdrawal application for UGX {amount} has been submitted and is pending approval.',
                withdrawal=withdrawal,
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Withdrawal application submitted successfully',
                'withdrawal_id': withdrawal.withdrawal_id,
                'amount': str(amount),
                'status': withdrawal.status,
                'purpose': purpose,
                'payment_method': payment_method,
                'application_date': withdrawal.application_date.isoformat()
            }, status=201)
        
        except json.JSONDecodeError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON format'
            }, status=400)
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)
    
    return JsonResponse({
        'success': False,
        'error': 'Only POST requests are allowed'
    }, status=405)


@login_required
def api_loan_status_approve(request, loan_id):
    """
    API endpoint to approve a loan
    Admin only
    Returns: JSON with status and loan details
    """
    if not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({
            'success': False,
            'error': 'Permission denied. Admin access required'
        }, status=403)
    
    try:
        loan = Loan.objects.get(pk=loan_id)
        
        if loan.status == 'approved':
            return JsonResponse({
                'success': False,
                'error': 'Loan is already approved'
            }, status=400)
        
        # Approve the loan
        loan.status = 'approved'
        loan.approval_date = timezone.now().date()
        loan.save()
        
        # Create notification for member
        Notification.objects.create(
            member=loan.member,
            notification_type='loan_approval',
            title='Your Loan Application has been Approved',
            message=f'Congratulations! Your loan application for UGX {loan.amount} at {loan.interest_rate}% interest for {loan.term_months} months has been APPROVED.',
            loan=loan,
            admin_user=request.user,
        )
        
        # Create notifications for other admins
        admin_members = Member.objects.filter(role__in=['admin', 'superadmin'])
        for admin_member in admin_members:
            if admin_member.user != request.user:
                Notification.objects.create(
                    member=admin_member,
                    notification_type='loan_approval',
                    title=f'Loan Approved: {loan.member.name}',
                    message=f'Admin {request.user.first_name or request.user.username} has APPROVED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).',
                    loan=loan,
                    admin_user=request.user,
                )
        
        return JsonResponse({
            'success': True,
            'message': 'Loan approved successfully',
            'loan_id': loan.loan_id,
            'status': loan.status,
            'approval_date': loan.approval_date.isoformat()
        }, status=200)
    
    except Loan.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Loan not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def api_loan_status_reject(request, loan_id):
    """
    API endpoint to reject a loan
    Admin only
    Returns: JSON with status and loan details
    """
    if not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({
            'success': False,
            'error': 'Permission denied. Admin access required'
        }, status=403)
    
    try:
        loan = Loan.objects.get(pk=loan_id)
        
        if loan.status == 'rejected':
            return JsonResponse({
                'success': False,
                'error': 'Loan is already rejected'
            }, status=400)
        
        # Reject the loan
        loan.status = 'rejected'
        loan.approval_date = timezone.now().date()
        loan.save()
        
        # Create notification for member
        Notification.objects.create(
            member=loan.member,
            notification_type='loan_rejection',
            title='Your Loan Application has been Rejected',
            message=f'Unfortunately, your loan application for UGX {loan.amount} has been REJECTED. Please contact admin for more details.',
            loan=loan,
            admin_user=request.user,
        )
        
        # Create notifications for other admins
        admin_members = Member.objects.filter(role__in=['admin', 'superadmin'])
        for admin_member in admin_members:
            if admin_member.user != request.user:
                Notification.objects.create(
                    member=admin_member,
                    notification_type='loan_rejection',
                    title=f'Loan Rejected: {loan.member.name}',
                    message=f'Admin {request.user.first_name or request.user.username} has REJECTED loan #{loan.loan_id} for {loan.member.name} (UGX {loan.amount}).',
                    loan=loan,
                    admin_user=request.user,
                )
        
        return JsonResponse({
            'success': True,
            'message': 'Loan rejected successfully',
            'loan_id': loan.loan_id,
            'status': loan.status,
            'approval_date': loan.approval_date.isoformat()
        }, status=200)
    
    except Loan.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Loan not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def api_loan_status(request, loan_id):
    """
    API endpoint to get loan status
    Returns: JSON with loan details
    """
    try:
        loan = Loan.objects.get(pk=loan_id)
        
        # Check permissions - only member or admin can view
        member = Member.objects.filter(user=request.user).first()
        if member and loan.member != member and not (request.user.is_staff or request.user.is_superuser):
            return JsonResponse({
                'success': False,
                'error': 'Permission denied'
            }, status=403)
        
        return JsonResponse({
            'success': True,
            'loan_id': loan.loan_id,
            'member_name': loan.member.name,
            'amount': str(loan.amount),
            'interest_rate': str(loan.interest_rate),
            'term_months': loan.term_months,
            'purpose': loan.purpose,
            'status': loan.status,
            'application_date': loan.application_date.isoformat(),
            'approval_date': loan.approval_date.isoformat() if loan.approval_date else None
        }, status=200)
    
    except Loan.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Loan not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)
